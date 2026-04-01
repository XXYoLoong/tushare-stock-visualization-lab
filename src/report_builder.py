from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from src.analysis import AnalysisBundle
from src.config import BASE_DIR, REPORT_DIR, Settings


def _set_document_style(document: Document) -> None:
    document.styles["Normal"].font.name = "Microsoft YaHei"
    document.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    document.styles["Normal"].font.size = Pt(11)


def _add_heading(document: Document, text: str, level: int = 1) -> None:
    heading = document.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _add_kv_table(document: Document, title: str, items: list[tuple[str, str]]) -> None:
    document.add_paragraph(title)
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    header = table.rows[0].cells
    header[0].text = "项目"
    header[1].text = "内容"
    for key, value in items:
        row = table.add_row().cells
        row[0].text = str(key)
        row[1].text = str(value)


def _add_image(document: Document, image_path: Path, caption: str) -> None:
    document.add_picture(str(image_path), width=Inches(6.4))
    paragraph = document.add_paragraph(caption)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_code_block(document: Document, file_path: Path) -> None:
    _add_heading(document, f"代码附录：{file_path.relative_to(BASE_DIR)}", level=2)
    paragraph = document.add_paragraph()
    run = paragraph.add_run(file_path.read_text(encoding="utf-8"))
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(8)


class ReportBuilder:
    def __init__(self, settings: Settings, bundle: AnalysisBundle) -> None:
        self.settings = settings
        self.bundle = bundle

    def build(self) -> Path:
        document = Document()
        _set_document_style(document)
        self._build_cover(document)
        self._build_overview(document)
        self._build_task1(document)
        self._build_task2(document)
        self._build_task3(document)
        self._build_code_appendix(document)

        path = REPORT_DIR / "Tushare股票数据可视化分析实验报告.docx"
        document.save(path)
        return path

    def _build_cover(self, document: Document) -> None:
        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run(self.settings.project_title)
        title_run.bold = True
        title_run.font.size = Pt(20)

        subtitle = document.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.add_run("基于 Tushare 的数据获取、分析、可视化与实验报告")

        meta = document.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta.add_run(f"生成时间：{datetime.now():%Y-%m-%d %H:%M:%S}")
        document.add_page_break()

    def _build_overview(self, document: Document) -> None:
        _add_heading(document, "一、实验说明", level=1)
        document.add_paragraph(
            "本实验依据用户提供的作业细则完成。由于原始语雀页面外部访问返回 403，"
            "本报告以用户消息中给出的题目要求为准。数据来源为 Tushare Pro 日线行情接口。"
        )
        _add_kv_table(
            document,
            "实验对象概览",
            [
                ("题目一", f"{self.bundle.task1.stock.name}（{self.bundle.task1.stock.ts_code}）2025 年日线行情"),
                ("题目二", "中国平安、比亚迪、贵州茅台 2025 年收益率对比"),
                ("题目三", f"{self.bundle.task3.stock.name} 最近 1 年均线趋势与交易信号分析"),
            ],
        )

    def _build_task1(self, document: Document) -> None:
        task1 = self.bundle.task1
        _add_heading(document, "二、题目一：单只股票日线行情分析与可视化", level=1)
        document.add_paragraph(
            "对研究股票按交易日期升序排序后进行缺失值和重复值检查，并统计收盘价与成交量特征。"
        )
        _add_kv_table(
            document,
            "清洗与统计结果",
            [
                ("股票", f"{task1.stock.name}（{task1.stock.ts_code}）"),
                ("记录数", str(task1.total_rows)),
                ("重复值数量", str(task1.duplicate_rows)),
                ("缺失值单元格数量", str(task1.missing_cells)),
                ("收盘价均值", f"{task1.closing_price_mean:.2f}"),
                ("收盘价标准差", f"{task1.closing_price_std:.2f}"),
                ("成交量均值", f"{task1.volume_mean:,.2f}"),
                ("最高收盘价", f"{task1.highest_close:.2f}"),
                ("最低收盘价", f"{task1.lowest_close:.2f}"),
            ],
        )
        _add_image(document, task1.volume_chart, "图 1 题目一每日成交量柱状图")
        _add_image(document, task1.monthly_chart, "图 2 题目一每月收盘价均值折线图")

    def _build_task2(self, document: Document) -> None:
        task2 = self.bundle.task2
        _add_heading(document, "三、题目二：多只股票收益率对比分析", level=1)
        document.add_paragraph(
            "选择金融、汽车新能源、食品饮料三类风格不同的股票，对 2025 年累计收益率进行比较，并展示其中一只股票的短中期均线变化。"
        )
        _add_kv_table(
            document,
            "最终累计收益率",
            [(name, f"{value:.2%}") for name, value in task2.final_returns.items()],
        )
        _add_image(document, task2.cumulative_return_bar_chart, "图 3 三只股票最终累计收益率柱状图")
        _add_image(document, task2.ma_chart, "图 4 任选股票收盘价与 5 日、20 日均线图")

    def _build_task3(self, document: Document) -> None:
        task3 = self.bundle.task3
        _add_heading(document, "四、题目三：趋势识别与简单交易信号分析", level=1)
        document.add_paragraph(
            "采用 5 日均线上穿 20 日均线作为买入信号，下穿 20 日均线作为卖出信号，"
            "并使用滞后一日持仓计算策略收益。"
        )
        _add_kv_table(
            document,
            "策略指标",
            [
                ("股票", f"{task3.stock.name}（{task3.stock.ts_code}）"),
                ("买入信号数", str(task3.buy_signals)),
                ("卖出信号数", str(task3.sell_signals)),
                ("股票累计收益率", f"{task3.stock_cumulative_return:.2%}"),
                ("策略累计收益率", f"{task3.strategy_cumulative_return:.2%}"),
                ("最大回撤", f"{task3.max_drawdown:.2%}"),
                ("完整交易次数", str(task3.completed_trades)),
                ("盈利交易次数", str(task3.profitable_trades)),
            ],
        )
        _add_image(document, task3.ma_chart, "图 5 收盘价与 5 日、20 日均线图")
        _add_image(document, task3.signal_chart, "图 6 买卖信号标注图")
        _add_image(document, task3.cumulative_return_chart, "图 7 累计收益率变化图")

        document.add_paragraph("结果分析：")
        for paragraph_text in task3.analysis_text:
            document.add_paragraph(paragraph_text, style="List Bullet")

        document.add_section(WD_SECTION.NEW_PAGE)

    def _build_code_appendix(self, document: Document) -> None:
        _add_heading(document, "五、代码附录", level=1)
        document.add_paragraph("按照提交要求，以下附上本实验核心代码文件。")
        for relative_path in [
            "src/config.py",
            "src/analysis.py",
            "src/report_builder.py",
            "src/main.py",
        ]:
            _add_code_block(document, BASE_DIR / relative_path)
